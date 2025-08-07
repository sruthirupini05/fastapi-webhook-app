import os
import io
import requests
import PyPDF2
from docx import Document
from urllib.parse import urlparse
from fastapi import FastAPI, HTTPException, UploadFile, File
from pydantic import BaseModel, Field
from typing import List, Union
import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document as LangchainDocument
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

app = FastAPI(title="LLM-Powered Intelligent Query-Retrieval System")

# --- FRONTEND SERVING ENDPOINTS ---
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/", response_class=FileResponse)
async def read_root():
    return "static/index.html"
# ----------------------------------

# --- Environment Variables ---
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY environment variable is not set. Please set it before running the application.")

# Initialize Google GenAI client
genai.configure(api_key=GOOGLE_API_KEY)
llm_model = genai.GenerativeModel('gemini-1.5-flash')

# Initialize embeddings model
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# --- Pydantic Models for API Request and Response ---
class QuestionRequest(BaseModel):
    # MODIFICATION: The Pydantic model already handles empty lists gracefully.
    documents: List[str] = Field(default_factory=list, description="List of document URLs to process.")
    questions: List[str] = Field(..., min_length=1, description="List of questions to answer.")

class QuestionResponse(BaseModel):
    answers: List[str]

# --- Helper Functions for Document Handling ---
def parse_document_content(content: bytes, filename: str):
    """Extracts text from PDF or DOCX content."""
    filename = filename.lower()
    if filename.endswith('.pdf'):
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = "".join([page.extract_text() or "" for page in reader.pages])
        return text
    elif filename.endswith('.docx'):
        doc = Document(io.BytesIO(content))
        text = "".join([para.text + "\n" for para in doc.paragraphs])
        return text
    else:
        raise ValueError("Unsupported document format. Only PDF and DOCX are supported.")

async def download_and_parse_document(document_url: str):
    """Downloads a document from a URL and extracts text."""
    print(f"DEBUG: Attempting to download document from URL: {document_url}")
    try:
        response = requests.get(document_url)
        response.raise_for_status()
        print(f"DEBUG: Successfully downloaded document from {document_url}. Status Code: {response.status_code}")
        content = response.content
        parsed_url = urlparse(document_url)
        filename = os.path.basename(parsed_url.path)
        
        if filename.endswith('.docx'):
            print(f"DEBUG: Attempting to parse DOCX file from {document_url}")
        
        return parse_document_content(content, filename)
    
    except requests.exceptions.RequestException as e:
        print(f"ERROR: Failed to download document from {document_url}. Specific error: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to download document from {document_url}: {e}")
    except Exception as e:
        print(f"ERROR: Failed to parse document from {document_url}. Specific error: {type(e).__name__} - {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse document from {document_url}: {type(e).__name__} - {e}")

# --- Main API Endpoint ---
# MODIFICATION: Removed duplicate endpoint
@app.post("/hackrx/run", response_model=QuestionResponse)
async def run_query_retrieval(
    request: QuestionRequest
):
    """
    Main endpoint to process documents, retrieve relevant clauses,
    and answer questions using an LLM. Documents must be provided as
    URLs in a JSON body.
    """
    # MODIFICATION: The check for an empty documents list has been removed.
    # The Pydantic model handles the data type, and the rest of the code
    # can gracefully handle an empty list.
    
    # MODIFICATION: Added a check for questions, as the logic requires it.
    if not request.questions:
        raise HTTPException(status_code=400, detail="At least one question must be provided.")
    
    all_docs_text = ""
    
    # Process documents from URLs
    for doc_url in request.documents:
        try:
            all_docs_text += await download_and_parse_document(doc_url)
            all_docs_text += "\n\n---END-OF-DOCUMENT---\n\n"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error processing document from {doc_url}: {e}")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, 
        chunk_overlap=200, 
        length_function=len
    )
    docs = [LangchainDocument(page_content=chunk) for chunk in text_splitter.split_text(all_docs_text)]
    
    # MODIFICATION: Handle the case where no documents are provided and text is empty
    if not all_docs_text.strip():
        # This allows the endpoint to still process and answer questions without a document,
        # though the LLM will state it cannot find the answer in the context.
        docs = [LangchainDocument(page_content="No documents provided.")]
    
    try:
        vectorstore = FAISS.from_documents(docs, embeddings)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create vector store: {type(e).__name__} - {e}")
    
    answers = []
    
    for question in request.questions:
        retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
        relevant_docs = retriever.invoke(question)
        retrieved_context = " ".join([doc.page_content for doc in relevant_docs])
        
        prompt = f"""
        You are an expert in policy documents. Your task is to answer a user's question
        based on the provided context from a policy document. If the answer is not
        explicitly stated in the context, state that you cannot find the information.

        Question: {question}

        Context:
        {retrieved_context}

        Answer:
        """
        
        try:
            response = llm_model.generate_content(prompt)
            answer_text = response.text.strip()
            answers.append(answer_text)
        except Exception as e:
            answers.append(f"An error occurred while generating the answer: {type(e).__name__} - {e}")
            
    return QuestionResponse(answers=answers)